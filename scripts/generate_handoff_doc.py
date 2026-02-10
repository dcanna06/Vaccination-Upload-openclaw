"""Generate developer handoff Word document for AIR Bulk Vaccination Upload System."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
    return table


def generate():
    doc = Document()

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("AIR Bulk Vaccination Upload System", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Developer Handoff & Architecture Guide")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version: 1.2.0\n").font.size = Pt(12)
    meta.add_run("Date: February 2026\n").font.size = Pt(12)
    meta.add_run("Classification: Internal — Developer Reference\n").font.size = Pt(12)

    doc.add_page_break()

    # Table of Contents placeholder
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Project Overview",
        "2. Technology Stack",
        "3. Architecture Overview",
        "4. Backend Structure",
        "5. Frontend Structure",
        "6. PRODA B2B Authentication",
        "7. AIR API Integration (All 16 APIs)",
        "8. Validation Rules",
        "9. Excel Template & Processing Pipeline",
        "10. Bulk Immunisation History Feature",
        "11. Testing Strategy",
        "12. Environment Configuration",
        "13. Deployment Guide",
        "14. NOI Certification",
        "15. Key Contacts",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # ================================================================
    # 1. Project Overview
    # ================================================================
    add_heading(doc, "1. Project Overview", level=1)
    doc.add_paragraph(
        "The AIR Bulk Vaccination Upload System is a web application that enables "
        "non-technical pharmacy and healthcare staff to upload Excel files containing "
        "vaccination records, validate them against AIR (Australian Immunisation Register) "
        "business rules, and submit them to the AIR via Services Australia REST Web Services."
    )
    doc.add_paragraph(
        "The system implements all 16 mandatory AIR APIs required for NOI (Notice of Integration) "
        "certification, including record encounter, individual management, exemptions, "
        "indicators, and reference data."
    )

    add_heading(doc, "Regulatory Context", level=2)
    doc.add_paragraph("The system is governed by:", style="List Bullet")
    regulations = [
        "Australian Immunisation Register Act 2015",
        "Privacy Act 1988 and Australian Privacy Principles",
        "My Health Records Act 2012",
        "Healthcare Identifiers Act 2010",
    ]
    for reg in regulations:
        doc.add_paragraph(reg, style="List Bullet 2")

    add_heading(doc, "Key Capabilities", level=2)
    capabilities = [
        "Bulk upload vaccination records from Excel files",
        "Server-side validation against AIR business rules (Medicare check digit, provider validation, date rules)",
        "Grouping of records into encounters and episodes (max 10 encounters, 5 episodes each)",
        "PRODA B2B OAuth 2.0 authentication with in-memory token caching",
        "Record encounter submission with confirmation flow (AIR-W-1004, AIR-W-1008)",
        "Individual management — identify, history, statement, exemptions, indicators",
        "Bulk immunisation history request — identify patients and fetch histories in batch",
        "Multi-location support with per-location Minor IDs",
        "Excel report download (4-sheet report with summary, history, vaccines due, errors)",
        "All 16 AIR APIs implemented for NOI certification",
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style="List Bullet")

    # ================================================================
    # 2. Technology Stack
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "2. Technology Stack", level=1)

    add_table(doc, ["Layer", "Technology", "Notes"], [
        ["Backend", "Python 3.12+ / FastAPI", "Async, type-hinted, Pydantic models"],
        ["Frontend", "Next.js 14 (App Router), React 18+", "TypeScript strict mode"],
        ["Styling", "TailwindCSS", "Dark theme, slate/emerald palette"],
        ["State Management", "Zustand + React Query", "Zustand for local state"],
        ["Excel Parsing", "openpyxl (backend), SheetJS (frontend)", "Server-side validation"],
        ["Database", "PostgreSQL 16", "Via SQLAlchemy 2.0 + Alembic"],
        ["Cache / Sessions", "Redis 7", "TLS-only in production"],
        ["Auth (App Users)", "JWT (HS256) + HttpOnly cookies", "8hr session, 30min timeout"],
        ["Auth (AIR API)", "PRODA B2B OAuth 2.0 JWT (RS256)", "60min token, auto-refresh"],
        ["Password Hashing", "Argon2id", "Not bcrypt"],
        ["HTTP Client", "httpx (async)", "For AIR API calls"],
        ["Logging", "structlog", "Structured JSON, no PII"],
        ["Testing", "pytest + Vitest + Playwright", "Unit, integration, E2E"],
        ["Infrastructure", "Azure (Australia East)", "DR in Australia Southeast"],
    ])

    # ================================================================
    # 3. Architecture Overview
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "3. Architecture Overview", level=1)

    add_heading(doc, "System Components", level=2)
    doc.add_paragraph(
        "The system follows a classic 3-tier architecture with the Next.js frontend communicating "
        "with the FastAPI backend, which integrates with Services Australia's AIR API via PRODA authentication."
    )

    doc.add_paragraph("Data Flow:", style="List Bullet")
    flow_steps = [
        "1. User uploads Excel file (.xlsx) through Next.js frontend",
        "2. Backend (FastAPI) parses Excel using openpyxl",
        "3. Server-side validation against AIR business rules",
        "4. Records grouped into encounters (max 10) and episodes (max 5 per encounter)",
        "5. PRODA B2B token acquired using JKS keystore (in-memory only)",
        "6. Batches submitted to AIR Record Encounter API v1.4",
        "7. Progress tracked in-memory with polling from frontend",
        "8. Results displayed with verbatim AIR messages",
        "9. Confirmation workflow for AIR-W-1004/AIR-W-1008 warnings",
        "10. Excel report download with 4 sheets",
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style="List Bullet 2")

    add_heading(doc, "Security Architecture", level=2)
    security_items = [
        "PRODA tokens stored in-memory only — never persisted to database or disk",
        "JKS keystore loaded from file or Base64 env var — never written to disk in production",
        "PII/PHI never logged — Medicare numbers masked in all log output",
        "structlog structured logging throughout — JSON format in production",
        "Rate limiting: 120 requests/minute per user",
        "Security headers middleware (X-Frame-Options, X-Content-Type-Options, etc.)",
        "CORS restricted to frontend URL only",
        "Argon2id password hashing (not bcrypt)",
    ]
    for item in security_items:
        doc.add_paragraph(item, style="List Bullet")

    # ================================================================
    # 4. Backend Structure
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "4. Backend Structure", level=1)

    add_heading(doc, "Directory Layout", level=2)
    doc.add_paragraph(
        "backend/\n"
        "  app/\n"
        "    main.py                    — FastAPI app factory, middleware, router registration\n"
        "    config.py                  — Pydantic Settings (env-driven configuration)\n"
        "    exceptions.py              — Custom exception classes\n"
        "    routers/                   — API endpoint handlers\n"
        "    services/                  — Business logic and external API clients\n"
        "    schemas/                   — Pydantic request/response models\n"
        "    middleware/                 — Auth, rate limiting, security headers, logging\n"
        "    models/                    — SQLAlchemy ORM models\n"
        "    utils/                     — Helper utilities\n"
        "  tests/\n"
        "    unit/                      — 531 unit tests\n"
        "    integration/               — NOI API suite, PRODA vendor tests\n"
        "  requirements.txt\n",
        style="No Spacing"
    )

    add_heading(doc, "Routers (API Endpoints)", level=2)
    add_table(doc, ["Router", "Prefix", "Purpose"], [
        ["health.py", "/health", "Health check endpoint"],
        ["auth.py", "/api/auth", "User authentication (register, login, logout)"],
        ["upload.py", "/api/upload", "Excel file upload and parsing"],
        ["template.py", "/api/template", "Download Excel template"],
        ["validate.py", "/api/validate", "Server-side record validation"],
        ["submit.py", "/api/submit", "Batch submission to AIR (background task)"],
        ["submission_results.py", "/api/submissions", "Submission history and results"],
        ["locations.py", "/api/locations", "Location CRUD with Minor ID management"],
        ["providers.py", "/api/providers", "Provider-location linking, HW027 status"],
        ["individuals.py", "/api/individuals", "Individual search, history, statement, vaccine trial"],
        ["encounters_update.py", "/api/encounters", "Update existing encounters (API #9)"],
        ["exemptions.py", "/api/exemptions", "Contraindication and natural immunity (APIs #5,6,10,11)"],
        ["indicators.py", "/api/indicators", "Vaccine indicators, indigenous status, catch-up (APIs #12-15)"],
        ["bulk_history.py", "/api/bulk-history", "Bulk immunisation history request (APIs #2,3)"],
    ])

    add_heading(doc, "Services (Business Logic)", level=2)
    add_table(doc, ["Service", "Purpose"], [
        ["proda_auth.py", "PRODA B2B token acquisition, JKS keystore loading, in-memory caching"],
        ["air_client.py", "AIR Record Encounter HTTP client with retry logic"],
        ["air_individual.py", "AIR Individual APIs — identify, history, statement, vaccine trial"],
        ["air_authorisation.py", "AIR Authorisation Access List (API #1)"],
        ["air_encounter_update.py", "AIR Update Encounter (API #9)"],
        ["air_exemptions.py", "AIR Contraindication and Natural Immunity (APIs #5,6,10,11)"],
        ["air_indicators.py", "AIR Vaccine Indicators, Indigenous Status, Catch-Up (APIs #12-15)"],
        ["validation_engine.py", "Server-side validation against AIR business rules"],
        ["excel_parser.py", "Excel file parsing with column mapping"],
        ["batch_grouping.py", "Group Excel rows into encounters and episodes"],
        ["submission_store.py", "Submission state management and persistence"],
        ["air_response_parser.py", "Parse AIR API responses (status codes, messages, claims)"],
        ["excel_template.py", "Generate Excel template for download"],
        ["location_manager.py", "Location and Minor ID business logic"],
        ["auth_service.py", "User authentication (Argon2id, JWT)"],
        ["air_resubmit.py", "Edit and resubmit corrected records"],
    ])

    # ================================================================
    # 5. Frontend Structure
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "5. Frontend Structure", level=1)

    add_heading(doc, "Pages", level=2)
    add_table(doc, ["Route", "Page", "Description"], [
        ["/login", "Login", "User authentication (email + password)"],
        ["/upload", "Upload", "Excel file upload with drag-and-drop"],
        ["/validate", "Validate", "Validation results with error display"],
        ["/submit", "Submit", "Submission progress monitor with pause/resume"],
        ["/history", "History", "Submission history list with filters"],
        ["/individuals", "Individuals", "Individual search (Medicare/IHI/demographics)"],
        ["/individuals/[id]", "Individual Detail", "Profile with history, statement, exemptions tabs"],
        ["/individuals/[id]/history", "History Details", "Immunisation history table"],
        ["/individuals/[id]/statement", "Statement", "Immunisation history statement download"],
        ["/individuals/[id]/exemptions", "Exemptions", "Contraindication and natural immunity"],
        ["/individuals/[id]/vaccinetrial", "Vaccine Trial", "Vaccine trial history"],
        ["/bulk-history", "Bulk History", "Bulk immunisation history request wizard"],
        ["/confirm", "Confirm", "Confirmation workflow for warnings"],
        ["/indicators", "Indicators", "Vaccine indicators and indigenous status"],
        ["/catchup", "Catch-Up", "Planned catch-up date scheduling"],
        ["/admin/locations", "Locations", "Location management (admin)"],
        ["/admin/providers", "Providers", "Provider-location linking (admin)"],
        ["/settings", "Settings", "Organisation settings and PRODA config"],
    ])

    add_heading(doc, "Key Components", level=2)
    components = [
        "Sidebar.tsx — Navigation sidebar with links to all pages",
        "FileUpload.tsx — Drag-and-drop file upload with validation",
        "ValidationResults.tsx — Error/warning display table",
        "SubmissionProgress.tsx — Real-time progress bar with polling",
        "SubmissionHistory.tsx — Submission history list with filtering",
        "ExcelTemplateDownload.tsx — Template download button",
    ]
    for comp in components:
        doc.add_paragraph(comp, style="List Bullet")

    add_heading(doc, "State Management", level=2)
    doc.add_paragraph(
        "The frontend uses Zustand for local state management. Key stores include "
        "uploadStore.ts (file state, parsed rows) and submissionStore.ts (batch progress). "
        "React Query may be used for server state in future iterations."
    )

    # ================================================================
    # 6. PRODA B2B Authentication
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "6. PRODA B2B Authentication", level=1)

    doc.add_paragraph(
        "The system uses PRODA (Provider Digital Access) B2B OAuth 2.0 for machine-to-machine "
        "authentication with the AIR API. This is implemented in backend/app/services/proda_auth.py."
    )

    add_heading(doc, "Token Acquisition Flow", level=2)
    token_steps = [
        "1. Load RSA private key from JKS keystore (SoapUI-generated, alias 'proda-alias')",
        "2. Build JWT assertion: iss=ORG_ID, sub=DEVICE_NAME, aud=proda URL, token.aud=PRODA audience, exp=10min",
        "3. Sign JWT with RS256 using private key (header: alg=RS256, kid=DEVICE_NAME, no typ)",
        "4. POST to PRODA token endpoint with grant_type=jwt-bearer, assertion, client_id",
        "5. Receive access_token (60min lifetime), key_expiry, device_expiry",
        "6. Cache token in-memory. Refresh at 50-minute mark (600s buffer).",
    ]
    for step in token_steps:
        doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "JWT Claims (Proven Working)", level=2)
    add_table(doc, ["Claim", "Value", "Notes"], [
        ["Header: alg", "RS256", "RSA with SHA-256"],
        ["Header: kid", "DEVICE_NAME", "e.g., DavidTestLaptop2"],
        ["Header: typ", "Not included", "Must set typ=False in PyJWT"],
        ["iss", "ORG_ID", "e.g., 2330016739 — NOT Minor ID"],
        ["sub", "DEVICE_NAME", "Same as kid header"],
        ["aud", "https://proda.humanservices.gov.au", "Fixed value"],
        ["token.aud", "https://proda.humanservices.gov.au", "Access token audience"],
        ["exp", "now + 600", "10 minutes"],
        ["iat", "now", "Unix timestamp"],
    ])

    add_heading(doc, "Critical Notes", level=2)
    critical_notes = [
        "PRODA tokens MUST be held in-memory only — never persist to database or disk",
        "JKS MUST be generated by SoapUI's ActivateDevice test suite — Python-generated keys are NOT compatible",
        "JWT iss is Org RA number, NOT Minor ID (previous versions had this wrong)",
        "JWT aud is always https://proda.humanservices.gov.au, NOT MCOL URL",
        "No jti claim — SoapUI doesn't include it and PRODA doesn't require it",
        "client_id is a POST form parameter, NOT a JWT claim",
        "Token endpoint differs between vendor and production environments",
    ]
    for note in critical_notes:
        doc.add_paragraph(note, style="List Bullet")

    # ================================================================
    # 7. AIR API Integration
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "7. AIR API Integration (All 16 APIs)", level=1)

    doc.add_paragraph(
        "Per AIR Developer Guide V3.0.8 Section 6.3, all 16 AIR Web Services must be implemented "
        "for NOI certification. The system implements all 16 APIs."
    )

    add_heading(doc, "Complete API Inventory", level=2)
    add_table(doc, ["#", "API Name", "Method", "Path", "Service File"], [
        ["1", "Authorisation Access List", "POST", "/air/immunisation/v1/authorisation/accesslist", "air_authorisation.py"],
        ["2", "Identify Individual", "POST", "/air/immunisation/v1.1/individual/details", "air_individual.py"],
        ["3", "Immunisation History Details", "POST", "/air/immunisation/v1.3/individual/history/details", "air_individual.py"],
        ["4", "Immunisation History Statement", "POST", "/air/immunisation/v1/individual/history/statement", "air_individual.py"],
        ["5", "Contraindication History", "POST", "/air/immunisation/v1/individual/contraindication/history", "air_exemptions.py"],
        ["6", "Record Contraindication", "POST", "/air/immunisation/v1/individual/contraindication/record", "air_exemptions.py"],
        ["7", "Vaccine Trial History", "POST", "/air/immunisation/v1/individual/vaccinetrial/history", "air_individual.py"],
        ["8", "Record Encounter", "POST", "/air/immunisation/v1.4/encounters/record", "air_client.py"],
        ["9", "Update Encounter", "POST", "/air/immunisation/v1.3/encounters/update", "air_encounter_update.py"],
        ["10", "Natural Immunity History", "POST", "/air/immunisation/v1/individual/naturalimmunity/history", "air_exemptions.py"],
        ["11", "Record Natural Immunity", "POST", "/air/immunisation/v1/individual/naturalimmunity/record", "air_exemptions.py"],
        ["12", "Add Vaccine Indicator", "POST", "/air/immunisation/v1/individual/vaccineindicator/add", "air_indicators.py"],
        ["13", "Remove Vaccine Indicator", "POST", "/air/immunisation/v1/individual/vaccineindicator/remove", "air_indicators.py"],
        ["14", "Update Indigenous Status", "POST", "/air/immunisation/v1/individual/indigenousstatus/update", "air_indicators.py"],
        ["15", "Planned Catch-Up Date", "POST", "/air/immunisation/v1.1/schedule/catchup", "air_indicators.py"],
        ["16", "Reference Data", "GET", "/air/immunisation/v1/refdata/*", "Direct httpx calls"],
    ])

    add_heading(doc, "Required HTTP Headers", level=2)
    doc.add_paragraph(
        "Every AIR API request MUST include all 11 mandatory headers. Missing any header returns 401 or 400."
    )
    add_table(doc, ["Header", "Format", "Notes"], [
        ["Authorization", "Bearer {token}", "PRODA B2B access token"],
        ["X-IBM-Client-Id", "API key", "From Developer Portal"],
        ["Content-Type", "application/json", "Always JSON"],
        ["Accept", "application/json", "Always JSON"],
        ["dhs-messageId", "urn:uuid:{UUID}", "Unique per request"],
        ["dhs-correlationId", "urn:uuid:{UUID}", "Same for related requests in a batch"],
        ["dhs-auditId", "Minor ID value", "Per-location Minor ID"],
        ["dhs-auditIdType", "Minor Id", "Literal string"],
        ["dhs-subjectId", "ddMMyyyy", "Patient DOB (e.g., 18102005)"],
        ["dhs-subjectIdType", "Date of Birth", "Literal string"],
        ["dhs-productId", "EM Bulk Vaccination Upload V1.2", "Software name + version"],
    ])

    add_heading(doc, "Response Handling", level=2)
    add_table(doc, ["Code", "Type", "Meaning", "Action"], [
        ["AIR-I-1007", "Info", "All encounters successfully recorded", "Mark as SUCCESS"],
        ["AIR-I-1000", "Info", "Individual encounter recorded", "Mark encounter as SUCCESS"],
        ["AIR-W-1004", "Warning", "Individual not found on AIR", "Prompt user to confirm"],
        ["AIR-W-1008", "Warning", "Some encounters not recorded", "Parse per-encounter results"],
        ["AIR-W-1001", "Warning", "Encounter NOT recorded", "Display warning, allow confirm"],
        ["AIR-E-1005", "Error", "Validation errors", "Display errors to user"],
        ["AIR-E-1006", "Error", "System error", "Retry with exponential backoff (max 3)"],
        ["AIR-E-1046", "Error", "Encounters not confirmable", "Correct and resubmit"],
    ])

    doc.add_paragraph(
        "CRITICAL: Error messages MUST be displayed to the end user exactly as supplied by "
        "Services Australia — not truncated, transformed, or modified in any way."
    )

    add_heading(doc, "Individual Identifier Pattern", level=2)
    doc.add_paragraph(
        "APIs #2-7 and #9-14 follow a two-step pattern:\n"
        "Step 1: Call Identify Individual (API #2) to get individualIdentifier (opaque, max 128 chars)\n"
        "Step 2: Pass individualIdentifier to subsequent calls\n\n"
        "Exceptions (do NOT use individualIdentifier): Record Encounter (#8), "
        "Planned Catch-Up Date (#15), Reference Data (#16)."
    )

    # ================================================================
    # 8. Validation Rules
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "8. Validation Rules", level=1)

    add_heading(doc, "Patient Identification (one required)", level=2)
    add_table(doc, ["Method", "Required Fields"], [
        ["Medicare", "medicareCardNumber (10 digits, check digit) + medicareIRN (1-9) + DOB + Gender"],
        ["IHI", "ihiNumber (16 digits, no Luhn) + DOB + Gender"],
        ["Demographics", "firstName + lastName + DOB + Gender + postCode"],
    ])

    add_heading(doc, "Field Validation", level=2)
    add_table(doc, ["Field", "Valid Values", "Notes"], [
        ["Gender", "M, F, X", "I and U are NOT valid in backend"],
        ["Route of Administration", "PO, SC, ID, IM, NS", "OR, IN, NAS are NOT valid in backend"],
        ["Vaccine Type", "NIP, OTH", "AEN is NOT valid in backend"],
        ["Vaccine Dose", "1-20 or B", "B = booster"],
        ["Medicare Card Number", "10 digits", "Check digit algorithm in Appendix A"],
        ["IHI Number", "16 digits", "Format only — NO Luhn check"],
        ["Provider Number", "8 characters", "Medicare or AIR format with check digit"],
        ["Date of Birth", "yyyy-MM-dd (API), ddMMyyyy (header)", "Not future, not >130 years ago"],
        ["Date of Service", "yyyy-MM-dd (API), ddMMyyyy (header)", "After DOB, not future"],
    ])

    add_heading(doc, "Medicare Check Digit Algorithm", level=2)
    doc.add_paragraph(
        "weighted_sum = (d1*1) + (d2*3) + (d3*7) + (d4*9) + (d5*1) + (d6*3) + (d7*7) + (d8*9)\n"
        "check_digit = weighted_sum % 10\n"
        "check_digit must equal digit 9\n"
        "digit 10 (issue number) must not be 0"
    )

    # ================================================================
    # 9. Excel Template & Processing Pipeline
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "9. Excel Template & Processing Pipeline", level=1)

    add_heading(doc, "Excel Columns", level=2)
    add_table(doc, ["Col", "Header", "Required", "Format"], [
        ["A", "Medicare Card Number", "Conditional", "10 digits"],
        ["B", "Medicare IRN", "Conditional", "1-9"],
        ["C", "IHI Number", "No", "16 digits"],
        ["D", "First Name", "Conditional", "Max 40 chars"],
        ["E", "Last Name", "Conditional", "Max 40 chars"],
        ["F", "Date of Birth", "Yes", "DD/MM/YYYY"],
        ["G", "Gender", "Yes", "M/F/X"],
        ["H", "Postcode", "Conditional", "4 digits"],
        ["I", "Date of Service", "Yes", "DD/MM/YYYY"],
        ["J", "Vaccine Code", "Yes", "1-6 chars (from ref data)"],
        ["K", "Vaccine Dose", "Yes", "1-20 or B"],
        ["L", "Vaccine Batch", "Conditional", "1-15 chars"],
        ["M", "Vaccine Type", "Conditional", "NIP/OTH"],
        ["N", "Route of Administration", "Conditional", "PO/SC/ID/IM/NS"],
        ["O", "Administered Overseas", "No", "TRUE/FALSE"],
        ["P", "Country Code", "Conditional", "3-char ISO 3166-1"],
        ["Q", "Immunising Provider Number", "Yes", "6-8 chars"],
        ["R", "School ID", "No", "Valid format"],
        ["S", "Antenatal Indicator", "No", "TRUE/FALSE"],
    ])

    add_heading(doc, "Processing Pipeline", level=2)
    pipeline = [
        "1. Parse all rows from Excel (openpyxl)",
        "2. Validate each row against AIR business rules",
        "3. Group by individual identity (Medicare+IRN or IHI or Name+DOB+Gender+Postcode)",
        "4. Within each group, sub-group by dateOfService — each becomes an encounter",
        "5. Within each encounter, each vaccine row becomes an episode (max 5 per encounter)",
        "6. Chunk into API requests of max 10 encounters each",
        "7. Submit sequentially (not parallel, to respect rate limits)",
        "8. Track progress in-memory, poll from frontend",
    ]
    for step in pipeline:
        doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "Performance", level=2)
    doc.add_paragraph(
        "500-row parse, validate, and group pipeline completes in under 1 second "
        "with less than 100MB peak memory usage."
    )

    # ================================================================
    # 10. Bulk Immunisation History Feature
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "10. Bulk Immunisation History Feature", level=1)

    doc.add_paragraph(
        "Allows users to upload an Excel file with patient identification details, "
        "validate them, then bulk-fetch immunisation history from AIR for all patients. "
        "Implemented in backend/app/routers/bulk_history.py and "
        "frontend/app/(dashboard)/bulk-history/page.tsx."
    )

    add_heading(doc, "User Flow", level=2)
    flow = [
        "1. Upload: User uploads .xlsx with patient ID columns",
        "2. Validate & Edit: System validates DOB, gender, Medicare/IHI. User can edit/skip invalid rows.",
        "3. Process: Background task calls Identify Individual (API #2) then Get History (API #3) for each patient",
        "4. Results: Summary cards + expandable per-patient history",
        "5. Download: 4-sheet Excel report (Summary, History, Vaccines Due, Errors)",
    ]
    for step in flow:
        doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "API Endpoints", level=2)
    add_table(doc, ["Method", "Path", "Description"], [
        ["POST", "/api/bulk-history/upload", "Upload and parse Excel file"],
        ["POST", "/api/bulk-history/validate", "Validate identification fields"],
        ["POST", "/api/bulk-history/process", "Start background processing"],
        ["GET", "/api/bulk-history/{id}/progress", "Poll processing progress"],
        ["GET", "/api/bulk-history/{id}/results", "Get completed results"],
        ["GET", "/api/bulk-history/{id}/download", "Download Excel report"],
    ])

    # ================================================================
    # 11. Testing Strategy
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "11. Testing Strategy", level=1)

    add_heading(doc, "Test Summary", level=2)
    add_table(doc, ["Category", "Count", "Framework", "Location"], [
        ["Backend unit tests", "531", "pytest + pytest-asyncio", "backend/tests/unit/"],
        ["Backend NOI integration", "23", "pytest (vendor env)", "backend/tests/integration/test_noi_api_suite.py"],
        ["Backend PRODA integration", "3", "pytest (vendor env)", "backend/tests/integration/test_proda_vendor.py"],
        ["Frontend unit tests", "150", "Vitest + React Testing Library", "frontend/__tests__/"],
        ["Playwright E2E tests", "73", "Playwright", "frontend/e2e/"],
        ["Performance tests", "8", "pytest (included in unit count)", "backend/tests/unit/test_performance.py"],
    ])

    add_heading(doc, "NOI Integration Tests", level=2)
    doc.add_paragraph(
        "The NOI test suite (test_noi_api_suite.py) tests all 16 APIs against the vendor environment. "
        "It uses shared state (individualIdentifier) across tests since API #2 must run before APIs #3-14."
    )
    noi_tests = [
        "API #1: Authorisation Access List — verify provider access",
        "API #2: Identify Individual — Medicare and IHI identification",
        "API #3: Immunisation History Details",
        "API #4: Immunisation History Statement",
        "API #5: Contraindication History",
        "API #6: Record Contraindication",
        "API #7: Vaccine Trial History",
        "API #8: Record Encounter",
        "API #9: Update Encounter",
        "API #10: Natural Immunity History",
        "API #11: Record Natural Immunity",
        "API #12: Add Vaccine Indicator",
        "API #13: Remove Vaccine Indicator",
        "API #14: Update Indigenous Status",
        "API #15: Planned Catch-Up Date",
        "API #16: Reference Data (vaccine + route endpoints)",
        "Workflow 1-5: All 5 record encounter use cases from TECH.SIS.AIR.02 Section 6",
    ]
    for test in noi_tests:
        doc.add_paragraph(test, style="List Bullet")

    add_heading(doc, "Running Tests", level=2)
    doc.add_paragraph(
        "Backend unit tests:\n"
        "  cd backend && .venv/bin/python -m pytest tests/unit/ -v\n\n"
        "Backend integration tests (requires PRODA credentials):\n"
        "  cd backend && .venv/bin/python -m pytest tests/integration/ -m integration -v\n\n"
        "Frontend unit tests:\n"
        "  cd frontend && npx vitest run\n\n"
        "Playwright E2E tests (requires running servers):\n"
        "  cd frontend && npx playwright test",
        style="No Spacing"
    )

    # ================================================================
    # 12. Environment Configuration
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "12. Environment Configuration", level=1)

    add_heading(doc, "Required Environment Variables", level=2)
    add_table(doc, ["Variable", "Description", "Example"], [
        ["APP_ENV", "Environment (vendor/production)", "vendor"],
        ["PRODA_ORG_ID", "PRODA Organisation ID (RA number)", "2330016739"],
        ["PRODA_DEVICE_NAME", "Device name registered with PRODA", "DavidTestLaptop2"],
        ["PRODA_MINOR_ID", "Minor ID for location", "WRR00000"],
        ["PRODA_JKS_FILE_PATH", "Local path to JKS keystore", "/path/to/keystore.jks"],
        ["PRODA_JKS_BASE64", "Base64-encoded JKS (alternative)", "(base64 string)"],
        ["PRODA_JKS_PASSWORD", "JKS keystore password", "Pass-123"],
        ["PRODA_KEY_ALIAS", "Private key alias in JKS", "proda-alias"],
        ["PRODA_JWT_AUDIENCE", "JWT aud claim", "https://proda.humanservices.gov.au"],
        ["PRODA_CLIENT_ID", "PRODA client ID", "soape-testing-client-v2"],
        ["PRODA_ACCESS_TOKEN_AUDIENCE", "JWT token.aud claim", "https://proda.humanservices.gov.au"],
        ["AIR_CLIENT_ID", "X-IBM-Client-Id from Developer Portal", "(UUID)"],
        ["AIR_PRODUCT_ID", "dhs-productId header value", "EM Bulk Vaccination Upload V1.2"],
        ["DATABASE_URL", "PostgreSQL connection string", "postgresql+asyncpg://..."],
        ["FRONTEND_URL", "Frontend URL for CORS", "http://localhost:3000"],
    ])

    add_heading(doc, "Vendor vs Production", level=2)
    add_table(doc, ["Setting", "Vendor", "Production"], [
        ["APP_ENV", "vendor", "production"],
        ["Token Endpoint", "vnd.proda.humanservices.gov.au/...", "proda.humanservices.gov.au/..."],
        ["AIR Base URL", "test.healthclaiming.api...", "Provided after NOI"],
        ["Client ID", "soape-testing-client-v2", "Assigned by Developer Portal"],
    ])

    # ================================================================
    # 13. Deployment Guide
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "13. Deployment Guide", level=1)

    add_heading(doc, "Local Development", level=2)
    local_steps = [
        "1. Clone repository: git clone <repo-url>",
        "2. Start PostgreSQL and Redis: docker compose -f infrastructure/docker-compose.yml up -d",
        "3. Backend setup: cd backend && python -m venv .venv && source .venv/bin/activate && pip install -r requirements.txt",
        "4. Copy .env.example to .env and configure PRODA/AIR credentials",
        "5. Run backend: uvicorn app.main:app --reload --port 8000",
        "6. Frontend setup: cd frontend && npm install",
        "7. Run frontend: npm run dev",
        "8. Access app at http://localhost:3000",
    ]
    for step in local_steps:
        doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "Production Deployment (Azure)", level=2)
    prod_steps = [
        "Deploy backend as Azure Container App",
        "Deploy frontend as Azure Static Web App or Container App",
        "PostgreSQL 16 Flexible Server (Australia East)",
        "Redis Premium P1 (TLS-only, zone-redundant)",
        "Azure Key Vault (Premium, HSM) for secrets",
        "Store JKS keystore Base64 in Key Vault",
        "Store all credentials in Key Vault — never in env files",
        "Configure CORS to allow only production frontend URL",
        "Enable TLS 1.2+ on all services",
        "Blue-green deployment with approval gates",
    ]
    for step in prod_steps:
        doc.add_paragraph(step, style="List Bullet")

    # ================================================================
    # 14. NOI Certification
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "14. NOI Certification", level=1)

    doc.add_paragraph(
        "NOI (Notice of Integration) certification is required from Services Australia before "
        "production access. The system has all required components implemented."
    )

    add_heading(doc, "Required Documentation", level=2)
    noi_docs = [
        "Application Details Form (docs/APPLICATION_DETAILS_FORM.md) — software details, tech stack, API list",
        "Integration Test Plan — NOI test suite covering all 16 APIs + 5 workflow scenarios",
        "User Manual (docs/user-guide.md) — step-by-step instructions for pharmacy staff",
        "GUI Screenshots — all key screens (upload, validate, submit, results, confirm, individual management)",
        "Architecture Summary — data flow, security measures, error handling",
    ]
    for d in noi_docs:
        doc.add_paragraph(d, style="List Bullet")

    add_heading(doc, "Submission Process", level=2)
    submission_steps = [
        "1. Complete all integration testing in vendor environment",
        "2. Contact OTS team: itest@servicesaustralia.gov.au",
        "3. Submit documentation package",
        "4. OTS team reviews (2-4 weeks turnaround)",
        "5. Address any feedback",
        "6. Receive NOI approval certificate",
        "7. Receive production credentials and endpoint URLs",
    ]
    for step in submission_steps:
        doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "5 Mandatory Workflow Test Cases", level=2)
    add_table(doc, ["#", "Scenario", "Expected AIR Code", "Required Action"], [
        ["1", "Standard success flow", "AIR-I-1007", "Mark as SUCCESS"],
        ["2", "Request validation fails", "AIR-E-1005", "Display errors, allow correction"],
        ["3", "Individual not found", "AIR-W-1004", "Confirm with acceptAndConfirm=Y"],
        ["4", "Pended episodes", "AIR-W-1008", "Parse per-encounter, selective confirm"],
        ["5", "Non-confirmable errors", "AIR-E-1046", "Correct and resubmit"],
    ])

    # ================================================================
    # 15. Key Contacts
    # ================================================================
    doc.add_page_break()
    add_heading(doc, "15. Key Contacts", level=1)

    add_table(doc, ["Team", "Contact", "Purpose"], [
        ["Developer Liaison", "DeveloperLiaison@servicesaustralia.gov.au", "Registration, test data, production access"],
        ["OTS Technical Support", "1300 550 115", "Technical issues during development"],
        ["OTS Product Integration", "itest@servicesaustralia.gov.au", "NOI certification testing"],
        ["PRODA Support (Production)", "1800 700 199", "PRODA production issues"],
    ])

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("— End of Document —").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = "/home/david/Vaccination-Upload/docs/Developer_Handoff_Guide.docx"
    doc.save(output_path)
    print(f"Document saved to {output_path}")
    return output_path


if __name__ == "__main__":
    generate()
